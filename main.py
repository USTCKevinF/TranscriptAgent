import os
import re
import yaml
import tiktoken
from openai import OpenAI
from docx import Document
from docx.shared import RGBColor, Pt
from concurrent.futures import ThreadPoolExecutor


class InterviewAgent:
    def __init__(self, config: dict):
        """初始化 OpenAI API"""
        self.client = OpenAI(
            api_key=config['api_key'],
            base_url=config.get('base_url', "https://vip.apiyi.com/v1")
        )
        with open('prompts.yaml', 'r', encoding='utf-8') as file:
            self.prompts = yaml.safe_load(file)

        self.model = config.get('model', 'gpt-4o-mini')
        self.interviewee_name = config['interviewee_name']
        self.temperature = config.get('temperature', 0.7)
        self.revise_iteration = config.get('revise_iteration', 1)
        self.chunk_size = config.get('chunk_size', 2000)

        # 获取transcript_system_prompt
        self.refinement_system_prompt = self.prompts['refinement_system_prompt']

        self.refined_introduction = self.refine_introduction(config['interviewee_introduction'])

        # self.transcript_system_prompt = self.prompts['transcript_system_prompt'].format(
        #     interviewee_name=self.interviewee_name,
        #     interviewee_introduction=self.refined_introduction
        # )

        self.transcript_system_prompt = self.prompts['transcript_system_prompt_V2']

        # 是否对文本进行优化润色
        self.enable_polish = config.get('enable_polish', True)

    def revise_text(self, unrevised_text: str) -> str:
        """调用 OpenAI API 进行对话"""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": self.transcript_system_prompt},
                    {"role": "user", "content": unrevised_text}
                ],
                temperature=self.temperature,
            )
            return response.choices[0].message.content
        
        except Exception as e:
            print(f"API 调用出错: {str(e)}")
            return None
        
    def polish_text(self, unrevised_text: str) -> str:
        """对文本进行优化润色"""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": self.prompts['polish_system_prompt']},
                    {"role": "user", "content": unrevised_text}
                ],
                # temperature=1.7,
            )
            return response.choices[0].message.content
        
        except Exception as e:
            print(f"出错了: {str(e)}")
            return None
        
    def refine_introduction(self, interviewee_introduction: str) -> str:
        """优化被采访者的自我介绍"""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": self.prompts['refinement_system_prompt']},
                    {"role": "user", "content": interviewee_introduction}
                ],
                temperature=0.7,
            )
            return response.choices[0].message.content
        
        except Exception as e:
            print(f"API 调用出错: {str(e)}")
            return None

    def check_difference(self, unrevised_text: str, revised_text: str) -> str:
        """检查润色前后的文本差异"""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": self.prompts['check_difference_system_prompt']},
                    {"role": "user", "content": f"转写前文本: {unrevised_text}\n转写后文本: {revised_text}"}
                ],
                temperature=1.0,
            )
            return response.choices[0].message.content 
        
        except Exception as e:
            print(f"API 调用出错: {str(e)}")
            return None

    def supply_missing_information(self, unrevised_text: str, revised_text: str, difference_information: str) -> str:
        """补充润色前文本中遗漏的信息"""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": self.prompts['supply_missing_information_system_prompt']},
                    {"role": "user", "content": self.prompts['supply_missing_information_user_prompt'].format(
                        unrevised_text=unrevised_text,
                        revised_text=revised_text,
                        difference_information=difference_information
                    )}
                ],
                temperature=0.2,
            )
            return response.choices[0].message.content 
        
        except Exception as e:
            print(f"API 调用出错: {str(e)}")
            return None
    
    def iterative_process_text(self, unrevised_text:str) -> str:
        """处理文本的主流程"""
        try:
            # 第一次补充
            revised_text = self.revise_text(unrevised_text)
            print("第一次转写结果：", revised_text)
            if revised_text is None:
                return None
                
            # 多次迭代补充和检查
            for _ in range(self.revise_iteration):
                # 检查差异
                difference = self.check_difference(unrevised_text, revised_text)
                print(difference)
                if difference is None:
                    return revised_text
                    
                # 补充遗漏信息
                revised_text = self.supply_missing_information(
                    unrevised_text,
                    revised_text, 
                    difference
                )
                
                print("补充前文本：", unrevised_text, "\n补充后文本：", revised_text, "\n差异信息：", difference)
                if revised_text is None:
                    return None
            
            # 对补充后的文本进行润色
            if self.enable_polish:
                print("润色前文本：", revised_text)
                revised_text = self.polish_text(revised_text)
                print("润色结果：", revised_text)
                if revised_text is None:
                    return None
                
            return revised_text
        
            
        except Exception as e:
            print(f"处理文本出错: {str(e)}")
            return None

    def remove_time_from_string(self,text:str) -> str:
        # 使用正则表达式匹配时间格式并替换为空字符串
        cleaned_text = re.sub(r'\(\d{2}:\d{2}:\d{2}\)', '', text)
        return cleaned_text.strip()  # 去掉前后的空格

    def replace_speaker(self, text:str, interviewee_name:str) -> str:
        speaker, content = text.split(': ', 1)
        if speaker != interviewee_name:
            speaker = "蜗壳进阶联盟"
        
        return f"{speaker}: {content}"

    def read_file_to_list(self, file_path:str) -> list:
        with open(file_path, 'r', encoding='utf-8') as file:
            # 读取所有行并去掉换行符
            lines = [line.strip() for line in file.readlines() if line.strip()]
        return lines
    
    def convert_format(self,file_path:str) -> str:
        content_list = self.read_file_to_list(file_path)
        content_list = [self.remove_time_from_string(line) for line in content_list]
        content_list = [self.replace_speaker(line, self.interviewee_name) for line in content_list]
        content_list = self.merge_consecutive_speakers(content_list)

        return content_list
    
    def chunk_text(self, content_list:list) -> list:
        chunks = []
        single_chunk = ""
        for line in content_list:
            if len(single_chunk) + len(line) > self.chunk_size:
                chunks.append(single_chunk)
                single_chunk = ""
            single_chunk += line
        if single_chunk:
            chunks.append(single_chunk)

        return chunks
    
    def count_tokens(self, text):
        encoding = tiktoken.encoding_for_model(self.model)
        tokens = encoding.encode(text)
        return len(tokens)
    
    def merge_consecutive_speakers(self, dialogue_list:list) -> list:
        if not dialogue_list:
            return []
        merged = []
        current_speaker = None
        current_content = ""
        
        for line in dialogue_list:
            speaker, content = line.split(': ', 1)
            
            if speaker == current_speaker:
                current_content += ' ' + content
            else:
                # 如果是新的说话人，保存之前的内容并开始新的
                if current_speaker:
                    merged.append(f"{current_speaker}: {current_content}")
                current_speaker = speaker
                current_content = content
    
        if current_speaker:
            merged.append(f"{current_speaker}: {current_content}")
        paired_dialogue = []
        for i in range(0, len(merged)-1, 2):
            if i+1 < len(merged):
                paired_dialogue.append(merged[i] + '\n' + merged[i+1])
            else:
                paired_dialogue.append(merged[i] + "")
        
        return paired_dialogue
    
    def iterative_process_text_list(self, text_list:list) -> list:
        revised_text_list = []
        for text in text_list:
            revised_text = self.iterative_process_text(text)
            revised_text_list.append(revised_text)
        return revised_text_list
    
    def process_text_standalone(self, text):
        """独立的处理函数"""
        return self.iterative_process_text(text)
    
    def multiprocess_iterative_process_text_list(self, text_list: list) -> list:
        """使用线程池处理文本列表"""
        with ThreadPoolExecutor(max_workers=8) as executor:
            revised_text_list = list(executor.map(self.iterative_process_text, text_list))
        return revised_text_list
    
    def format_interview_text(self, content, output_path):

        # 创建新文档
        doc = Document()

        # 将英文冒号替换为中文冒号
        content = content.replace(':', '：')
        
        # 按段落分割
        paragraphs = content.strip().split('\n\n')
        
        for para in paragraphs:
            # 创建新段落
            p = doc.add_paragraph()
            
            # 设置段前距和段后距
            p.paragraph_format.space_before = Pt(8)  # 段前距设置为8磅
            p.paragraph_format.space_after = Pt(24)  # 段后距设置为24磅

            if '：' in para:
                # 分割说话人和内容
                speaker, text = para.split('：', 1)    
                # 如果是蜗壳进阶联盟说话
                if speaker == '蜗壳进阶联盟':
                    # 添加加粗的红色说话人名字
                    speaker_run = p.add_run(f'{speaker}：')
                    speaker_run.bold = True
                    speaker_run.font.color.rgb = RGBColor(171, 25, 66)
                    speaker_run.font.name = 'SimHei'  # 设置字体为SimHei
                    speaker_run.font.size = Pt(15)  # 设置字号为15磅
                    
                    # 添加红色的内容
                    content_run = p.add_run(text)
                    content_run.bold = True
                    content_run.font.color.rgb = RGBColor(171, 25, 66)
                    content_run.font.name = 'SimHei'  # 设置字体为SimHei
                    content_run.font.size = Pt(15)  # 设置字号为15磅

                else:
                    # 其他说话人只加粗名字
                    speaker_run = p.add_run(f'{speaker}：')
                    speaker_run.bold = True
                    speaker_run.font.name = 'SimHei'  # 设置字体为SimHei
                    speaker_run.font.size = Pt(15)  # 设置字号为15磅
                    
                    # 添加普通内容
                    content_run = p.add_run(text)
                    content_run.font.name = 'SimHei'  # 设置字体为SimHei
                    content_run.font.size = Pt(15)  # 设置字号为15磅
        
            else:
                # 处理没有说话人的段落
                content_run = p.add_run(para)
                content_run.font.name = 'SimHei'  # 设置字体为SimHei
                content_run.font.size = Pt(15)  # 设置字号为15磅

        # 保存文档
        doc.save(output_path)

    def revise(self, file_path:str, output_path:str):
        converted_format_text = self.convert_format(file_path=file_path)
        chunked_list = self.chunk_text(content_list=converted_format_text)
        revised_result = self.multiprocess_iterative_process_text_list(text_list=chunked_list)
        final_result = ''
        for single_result in revised_result:
            final_result += single_result + '\n\n'

        self.format_interview_text(final_result, output_path)

if __name__ == "__main__":
    with open('./config.yaml', 'r', encoding='utf-8') as file:
        config = yaml.safe_load(file)
    
    # 确保输出目录存在
    output_dir = os.path.dirname(config['output_file'])
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    Agent = InterviewAgent(config)
    Agent.revise(
        file_path=config['input_file'],
        output_path=config['output_file']
    )